import openai
import os
from flask import Flask, request, url_for, redirect, render_template,session
import pickle
import numpy as np
from flask_bootstrap import Bootstrap
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from docx import Document
from threading import Thread
import pandas as pd
import time
from datetime import datetime
from forex_python.converter import CurrencyCodes
from flask import send_from_directory
import json 
from werkzeug.utils import secure_filename

app = Flask(__name__)

app.config['SECRET_KEY'] = "Hackathon"
Bootstrap(app)

MODEL_PATH = "model.pkl"

openai.api_key = "<YOUR-OPEN-AI-KEY>"

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
RESULT_FOLDER = 'results'
app.config['RESULT_FOLDER'] = RESULT_FOLDER

def extract_text_from_docx(file_path):
    """Extract text from the uploaded Word document."""
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def clean_generated_code(generated_code):
    # Split the code into lines
    lines = generated_code.splitlines()
    
    # Filter only the lines that start with 'df' and remove leading whitespace
    filtered_lines = [line.lstrip() for line in lines if line.strip().startswith('df')]
    
    # Join the remaining lines back into a single string
    cleaned_code = "\n".join(filtered_lines)

    print(cleaned_code)
    
    return cleaned_code


def generate_code_from_natural_language(rules_text, column_names):
    """Generate Python code for data processing from the rules and column names."""
    columns_str = ", ".join(column_names)
    
    prompt = f"""
    The following are the column names in the dataset: {columns_str}
    Convert the following profiling rules into Python code for a classification model. 
    The code should use these column names to flag and assign a risk score based on the rules:
    {rules_text}
    time is in this format "25-02-2025" "DD-MM-YYYY"
    Please return **only the Python code** for applying the rules on the dataset and assigning risk scores.
    The code should be executable, use the column names directly, and only return the Python code to assign risk scores.
    any other explanation should be given in comments only starting with # only. DOnt give explanation 
    please use only variables which are in column_name. 
    please ont create any new function call
    if there is something in rules which is not in domain of the columns return the string 'Profiling Rules must be in the domain of the data'
    """
    
    response = openai.ChatCompletion.create(
        model="gpt-4",  # Or use gpt-3.5-turbo for better cost management
        messages=[{"role": "system", "content": "You are a helpful assistant that generates Python code."},
                  {"role": "user", "content": prompt}],
        max_tokens=1000,
        temperature=0.3
    )
    
    generated_code = response['choices'][0]['message']['content'].strip()
    
    cleaned_code = clean_generated_code(generated_code)
    
    return cleaned_code

def analyze_row_with_ai(rules_text, row_data):
    
    prompt = f"""
    The following is a single record in the dataset: {row_data}
    Using these rules tell what's wrong with the data and how to remediate this 
    {rules_text}. Just give whats the problem and how to remediate part.
    """
    
    response = openai.ChatCompletion.create(
        model="gpt-4",  # Or use gpt-3.5-turbo for better cost management
        messages=[{"role": "system", "content": "You are a helpful assistant that tells whats the problem with data according to rules."},
                  {"role": "user", "content": prompt}],
        max_tokens=1000,
        temperature=0.3
    )
    
    generated_res = response['choices'][0]['message']['content'].strip()    
    return generated_res

def process_csv_data(file_path, generated_code):
    """Process the CSV data using the generated code."""
    df = pd.read_csv(file_path)
    try:
        exec(generated_code)
    except Exception as e:
        return f"Error executing code: {str(e)}"
    return df

def process_files(docx_file_path, csv_file_path):
    """Extract rules from DOCX file, generate code, and process the CSV data."""
    rules_text = extract_text_from_docx(docx_file_path)
    data = pd.read_csv(csv_file_path)
    column_names = list(data.columns)
    
    # Generate Python code based on rules and column names
    generated_code = generate_code_from_natural_language(rules_text, column_names)
    
    # Debug: Print generated code and column names for review
    print("Column Names:")
    print(column_names)
    
    # Process CSV data using the generated code
    result_data = process_csv_data(csv_file_path, generated_code)
    return result_data

def process_in_background(docx_file_path, csv_file_path):
    result_data = process_files(docx_file_path, csv_file_path)
    print(result_data)
    if isinstance(result_data,str):
        return "Some Error Occured"
    else:
        csv_file_path_out = os.path.join(app.config['RESULT_FOLDER'], "processed_data.csv")
        result_data.to_csv(csv_file_path_out, index=False)
        return result_data, csv_file_path_out


@app.route("/test", methods=["GET", "POST"])
def test_page():
    if request.method == "POST":
        if "docx_file" not in request.files or "csv_file" not in request.files:
            return redirect(request.url)
        
        docx_file = request.files["docx_file"]
        csv_file = request.files["csv_file"]
        
        if docx_file.filename == "" or csv_file.filename == "":
            return redirect(request.url)

        if docx_file and docx_file.filename.endswith(".docx") and csv_file and csv_file.filename.endswith(".csv"):

            secure_docx_filename = secure_filename(docx_file.filename)
            secure_csv_filename = secure_filename(csv_file.filename)

            new_docx_filename = "rules.docx"  # You can customize this as you wish
            new_csv_filename = "data.csv"     # You can customize this as you wish

            # Get the file paths where they will be saved
            docx_file_path = os.path.join(app.config['UPLOAD_FOLDER'], new_docx_filename)
            csv_file_path = os.path.join(app.config['UPLOAD_FOLDER'], new_csv_filename)
            
            docx_file.save(docx_file_path)
            csv_file.save(csv_file_path)

            res = process_in_background(docx_file_path, csv_file_path)

            if isinstance(res,str):
                return render_template("error.html")
            else:
                return render_template("flagged.html", res = res, csv_file=process_csv_data)
    
    return render_template("test.html", message="Please upload the Word and CSV files.")



@app.route("/")
def main_page():
    return render_template("index.html")


@app.route('/resuls/<filename>')
def download_csv(filename):
    return send_from_directory(app.config['RESULT_FOLDER'], filename)

def get_flagged_data():
    # Load the processed CSV file
    file_path = 'results/processed_data.csv'
    df = pd.read_csv(file_path)

    # Find the last column dynamically
    last_column = df.columns[-1]

    # Filter rows where the last column's value is not zero
    flagged_data = df[df[last_column] != 0]

    # Convert flagged data to a list of dictionaries (for easy rendering in templates)
    return flagged_data.to_dict(orient='records')

@app.route('/view_flagged_data')
def view_flagged_data():
    # Fetch or generate flagged rows (rows with risk_score != 0)
    flagged_data = get_flagged_data()  # You need to implement this function to get the data
    
    # Render the flagged data page with the flagged rows
    return render_template('view_flagged_data.html', flagged_data=flagged_data)

def get_row_data(col_value):
    # Load the processed CSV file
    file_path = 'results/processed_data.csv'
    df = pd.read_csv(file_path)

    # Get the name of the first column
    first_column = df.columns[0]

    # Find the row where the value in the first column matches `col_value`
    matching_row = df[df[first_column] == col_value]

    # Return the row as a dictionary (if a match is found)
    if not matching_row.empty:
        return matching_row.iloc[0].to_dict()
    else:
        return None

def format_analysis_result(analysis_result):
    # Split the result by newlines to handle each issue/remediation
    issues = analysis_result.split("\n")
    
    # Initialize the formatted result with an empty string
    formatted_result = ""
    
    # Loop through each issue and wrap with <p> tags
    for issue in issues:
        if issue.strip():  # Skip empty strings
            formatted_result += f"<p>{issue.strip()}</p>"
    
    return formatted_result

@app.route('/row_analysis/<int:row_id>')
def row_analysis(row_id):
    # Fetch the specific row's data based on the row_id
    row_data = get_row_data(row_id)  # You need to implement this function to get the row data
    doc = Document("uploads/rules.docx")
    rules_text = ""
    for para in doc.paragraphs:
        rules_text += para.text + "\n"
    
    # Analyze the row using the AI model and profiling rules
    analysis_result = analyze_row_with_ai(rules_text,row_data)  # Implement this function
    
    format_analysis = format_analysis_result(analysis_result)
    print(analysis_result)
    formatted_row_data = json.dumps(row_data, indent=4)
    # Render the analysis result
    return render_template('row_analysis.html', row_data=formatted_row_data, analysis_result=format_analysis)

if __name__ == "__main__":
    app.run(debug=True)